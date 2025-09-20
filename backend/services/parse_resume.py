# backend/services/parse_resume.py
import pdfplumber
import re
from io import BytesIO
from docx import Document
from typing import Dict, List, Optional
from ..utils.extract_entities import extract_emails, extract_phones

# Comprehensive skills dictionary
TECHNICAL_SKILLS = {
    # Programming Languages
    'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'c', 'go', 'rust', 'swift', 
    'kotlin', 'php', 'ruby', 'scala', 'r', 'matlab', 'sql', 'html', 'css', 'sass', 'less',
    
    # Frameworks & Libraries
    'react', 'angular', 'vue', 'node.js', 'nodejs', 'express', 'django', 'flask', 'fastapi',
    'spring', 'spring boot', 'laravel', 'rails', 'asp.net', 'tensorflow', 'pytorch', 
    'scikit-learn', 'pandas', 'numpy', 'matplotlib', 'seaborn', 'opencv', 'keras',
    
    # Databases
    'mongodb', 'postgresql', 'mysql', 'redis', 'elasticsearch', 'cassandra', 'oracle', 
    'sqlite', 'dynamodb', 'neo4j', 'firebase', 'supabase',
    
    # Cloud & DevOps
    'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'jenkins', 'git', 
    'github', 'gitlab', 'terraform', 'ansible', 'linux', 'ubuntu', 'bash', 'ci/cd',
    'nginx', 'apache', 'helm', 'prometheus', 'grafana',
    
    # Data & Analytics
    'machine learning', 'deep learning', 'nlp', 'computer vision', 'data science',
    'big data', 'hadoop', 'spark', 'kafka', 'airflow', 'etl', 'data mining', 'statistics',
    
    # Mobile & Frontend
    'react native', 'flutter', 'ionic', 'xamarin', 'android', 'ios', 'swift ui',
    
    # Other Technologies
    'rest api', 'graphql', 'microservices', 'agile', 'scrum', 'jira', 'confluence',
    'elasticsearch', 'solr', 'rabbitmq', 'redis', 'memcached'
}

# Education keywords
EDUCATION_KEYWORDS = {
    'degrees': ['b.tech', 'btech', 'b.e', 'be', 'bachelor', 'b.sc', 'bsc', 'b.com', 'bcom',
                'm.tech', 'mtech', 'm.e', 'me', 'master', 'm.sc', 'msc', 'm.com', 'mcom',
                'mba', 'mca', 'phd', 'doctorate', 'diploma', 'associate'],
    'fields': ['computer science', 'information technology', 'software engineering', 
               'electrical', 'mechanical', 'civil', 'electronics', 'mathematics',
               'physics', 'chemistry', 'business', 'management', 'finance', 'marketing']
}

# Certification keywords
CERTIFICATION_KEYWORDS = [
    'aws certified', 'azure certified', 'google cloud certified', 'cisco certified',
    'microsoft certified', 'oracle certified', 'certified', 'certification',
    'pmp', 'scrum master', 'product owner', 'csm', 'comptia', 'cissp'
]

# Section headers for experience/projects
SECTION_HEADERS = {
    'experience': ['experience', 'work experience', 'employment', 'work history', 'professional experience'],
    'projects': ['projects', 'personal projects', 'academic projects', 'key projects'],
    'education': ['education', 'academic background', 'qualifications'],
    'skills': ['skills', 'technical skills', 'technologies', 'competencies'],
    'certifications': ['certifications', 'certificates', 'achievements', 'awards']
}

def clean_and_normalize_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""
    
    # Remove excessive whitespace and normalize line breaks
    text = re.sub(r'\n+', '\n', text)
    text = re.sub(r' +', ' ', text)
    
    # Remove common header/footer patterns
    header_footer_patterns = [
        r'Page \d+ of \d+',
        r'Confidential',
        r'Resume of .*',
        r'CV of .*',
        r'\d+/\d+/\d+',  # dates
    ]
    
    for pattern in header_footer_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    
    return text.strip()

def extract_name(text: str) -> Optional[str]:
    """Extract candidate name from resume text"""
    lines = text.split('\n')[:5]  # Check first 5 lines
    
    # Look for name patterns
    name_patterns = [
        r'^([A-Z][a-z]+ [A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s*$',  # "John Smith" or "John Doe Smith"
        r'^([A-Z][A-Z\s]+)$',  # "JOHN SMITH"
    ]
    
    for line in lines:
        line = line.strip()
        if len(line) > 50 or len(line) < 5:  # Skip very long or short lines
            continue
            
        for pattern in name_patterns:
            match = re.match(pattern, line)
            if match:
                name = match.group(1).strip()
                # Validate it's not a common header
                if not any(keyword in name.lower() for keyword in ['resume', 'cv', 'curriculum']):
                    return name
    
    return None

def extract_skills_from_text(text: str) -> List[str]:
    """Extract technical skills from resume text"""
    text_lower = text.lower()
    found_skills = []
    
    for skill in TECHNICAL_SKILLS:
        # Use word boundaries for exact matching
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, text_lower):
            found_skills.append(skill)
    
    return sorted(list(set(found_skills)))

def extract_education(text: str) -> List[Dict[str, str]]:
    """Extract education information"""
    education = []
    text_lower = text.lower()
    
    # Find degree mentions
    for degree in EDUCATION_KEYWORDS['degrees']:
        pattern = r'\b' + re.escape(degree) + r'\b.*?(?:\n|\.|\||$)'
        matches = re.finditer(pattern, text_lower)
        
        for match in matches:
            context = match.group(0).strip()
            
            # Try to extract field and institution
            field = None
            for field_keyword in EDUCATION_KEYWORDS['fields']:
                if field_keyword in context:
                    field = field_keyword
                    break
            
            # Extract year (4 digits)
            year_match = re.search(r'\b(19|20)\d{2}\b', context)
            year = year_match.group(0) if year_match else None
            
            education.append({
                'degree': degree.upper(),
                'field': field,
                'year': year,
                'context': context[:100]  # First 100 chars for context
            })
    
    return education

def extract_certifications(text: str) -> List[str]:
    """Extract certifications from resume"""
    certifications = []
    text_lower = text.lower()
    
    for cert_keyword in CERTIFICATION_KEYWORDS:
        pattern = r'\b' + re.escape(cert_keyword) + r'\b.*?(?:\n|\.|\||$)'
        matches = re.finditer(pattern, text_lower)
        
        for match in matches:
            cert_context = match.group(0).strip()
            certifications.append(cert_context[:100])  # Limit context length
    
    return list(set(certifications))

def extract_experience_years(text: str) -> Optional[int]:
    """Extract years of experience from resume"""
    # Patterns for experience mentions
    experience_patterns = [
        r'(\d+)\+?\s*years?\s*of\s*experience',
        r'(\d+)\+?\s*years?\s*experience',
        r'experience\s*:\s*(\d+)\+?\s*years?',
        r'(\d+)\+?\s*yrs?\s*experience'
    ]
    
    text_lower = text.lower()
    for pattern in experience_patterns:
        match = re.search(pattern, text_lower)
        if match:
            return int(match.group(1))
    
    return None

def detect_sections(text: str) -> Dict[str, List[str]]:
    """Detect and extract content from different resume sections"""
    sections = {}
    lines = text.split('\n')
    current_section = None
    current_content = []
    
    for line in lines:
        line_lower = line.strip().lower()
        
        # Check if line is a section header
        section_found = False
        for section_type, headers in SECTION_HEADERS.items():
            if any(header in line_lower for header in headers):
                # Save previous section
                if current_section and current_content:
                    sections[current_section] = current_content
                
                current_section = section_type
                current_content = []
                section_found = True
                break
        
        if not section_found and current_section:
            current_content.append(line.strip())
    
    # Save last section
    if current_section and current_content:
        sections[current_section] = current_content
    
    return sections

async def extract_text_from_file(file) -> str:
    """Enhanced text extraction with cleaning"""
    content = await file.read()
    name = file.filename.lower()

    if name.endswith(".pdf"):
        raw_text = extract_text_from_pdf_bytes(content)
    elif name.endswith(".docx"):
        raw_text = extract_text_from_docx_bytes(content)
    else:
        raw_text = content.decode("utf-8", errors="ignore")
    
    return clean_and_normalize_text(raw_text)

def extract_text_from_pdf_bytes(b: bytes) -> str:
    """Enhanced PDF text extraction"""
    text_blocks = []
    
    with pdfplumber.open(BytesIO(b)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                # Clean page text
                cleaned_page = clean_and_normalize_text(page_text)
                if cleaned_page:
                    text_blocks.append(cleaned_page)
    
    return "\n".join(text_blocks)

def extract_text_from_docx_bytes(b: bytes) -> str:
    """Enhanced DOCX text extraction"""
    doc = Document(BytesIO(b))
    
    text_blocks = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_blocks.append(paragraph.text.strip())
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_blocks.append(cell.text.strip())
    
    return clean_and_normalize_text("\n".join(text_blocks))

async def parse_resume_comprehensive(file) -> Dict:
    """
    Comprehensive resume parsing with structured output
    Returns detailed JSON with all extracted entities
    """
    # Extract raw text
    raw_text = await extract_text_from_file(file)
    
    if not raw_text:
        return {"error": "Could not extract text from file"}
    
    # Extract all entities
    name = extract_name(raw_text)
    emails = extract_emails(raw_text)
    phones = extract_phones(raw_text)
    skills = extract_skills_from_text(raw_text)
    education = extract_education(raw_text)
    certifications = extract_certifications(raw_text)
    experience_years = extract_experience_years(raw_text)
    sections = detect_sections(raw_text)
    
    # Calculate resume statistics
    stats = {
        'total_characters': len(raw_text),
        'total_words': len(raw_text.split()),
        'total_lines': raw_text.count('\n') + 1,
        'skills_count': len(skills),
        'education_count': len(education),
        'certifications_count': len(certifications)
    }
    
    return {
        'personal_info': {
            'name': name,
            'emails': emails,
            'phones': phones
        },
        'technical_profile': {
            'skills': skills,
            'experience_years': experience_years,
            'certifications': certifications
        },
        'education': education,
        'sections': sections,
        'statistics': stats,
        'raw_text': raw_text,
        'filename': file.filename
    }
